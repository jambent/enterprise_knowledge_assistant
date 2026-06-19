from pathlib import Path
from typing import Iterator
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


class MultiFormatLoader(BaseLoader):
    def __init__(self, path: str):
        self.path = Path(path)

    def lazy_load(self) -> Iterator[Document]:
        suffix = self.path.suffix.lower()

        if suffix == ".txt":
            text = self._load_txt()
        elif suffix == ".pdf":
            text = self._load_pdf()
        elif suffix == ".docx":
            text = self._load_docx()
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        yield Document(
            page_content=text,
            metadata={
                "source": str(self.path),
                "file_type": suffix
            }
        )


    def _load_txt(self) -> str:
        with open(self.path, "r", encoding="utf-8") as f:
            return f.read()

    def _load_pdf(self) -> str:
        reader = PdfReader(str(self.path))
        text = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text.append(extracted)
        return "\n".join(text)

    def _load_docx(self) -> str:
        doc = DocxDocument(self.path)
        return "\n".join(p.text for p in doc.paragraphs)

    
if __name__ == "__main__":
    loader = MultiFormatLoader("./input_files/policies/CP01_Enterprise_knowledge_Assistant_Capstone.docx")
    for doc in loader.lazy_load():
        print(doc.page_content[:500])